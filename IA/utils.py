import PyPDF2
from docx import Document
import csv
import pandas as pd
import mimetypes


def calculate_num_timesteps_audio(input_data, sample_rate, timestep_duration):
    # Cálculo do num_timesteps para dados de áudio
    # input_data: array com os dados de áudio
    # sample_rate: taxa de amostragem dos dados de áudio (amostras por segundo)
    # timestep_duration: duração desejada de cada timestep em segundos
    
    audio_length = len(input_data)  # Comprimento total dos dados de áudio
    timestep_samples = int(sample_rate * timestep_duration)  # Número de amostras por timestep
    num_timesteps = int(audio_length / timestep_samples)  # Cálculo do num_timesteps
    return num_timesteps


def calculate_num_timesteps_video(input_data, frames_per_second, timestep_duration):
    # Cálculo do num_timesteps para dados de vídeo
    # input_data: array com os frames de vídeo
    # frames_per_second: número de frames por segundo
    # timestep_duration: duração desejada de cada timestep em segundos
    
    num_frames = len(input_data)  # Número total de frames no vídeo
    timestep_frames = int(frames_per_second * timestep_duration)  # Número de frames por timestep
    num_timesteps = int(num_frames / timestep_frames)  # Cálculo do num_timesteps
    return num_timesteps


def calculate_num_timesteps_image(input_data, image_width, timestep_width):
    # Cálculo do num_timesteps para dados de imagem
    # input_data: array com os pixels da imagem
    # image_width: largura da imagem em pixels
    # timestep_width: largura desejada de cada timestep em pixels
    
    num_pixels = len(input_data)  # Número total de pixels na imagem
    timestep_pixels = int(timestep_width * image_width)  # Número de pixels por timestep
    num_timesteps = int(num_pixels / timestep_pixels)  # Cálculo do num_timesteps
    return num_timesteps


def calculate_num_timesteps_text(input_data, timestep_length):
    # Cálculo do num_timesteps para dados de texto
    # input_data: string com o texto
    # timestep_length: comprimento desejado de cada timestep em caracteres
    
    text_length = len(input_data)  # Comprimento total do texto
    num_timesteps = int(text_length / timestep_length)  # Cálculo do num_timesteps
    return num_timesteps


def calculate_num_timesteps_file(input_data, timestep_size):
    # Cálculo do num_timesteps para outros tipos de arquivo
    # input_data: array com os dados do arquivo
    # timestep_size: tamanho desejado de cada timestep em bytes
    
    file_size = len(input_data)  # Tamanho total do arquivo
    num_timesteps = int(file_size / timestep_size)  # Cálculo do num_timesteps
    return num_timesteps

def calculate_num_timesteps_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf = PyPDF2.PdfFileReader(file)
        num_pages = pdf.getNumPages()
        return num_pages


def calculate_num_timesteps_csv(file_path):
    with open(file_path, 'r') as file:
        csv_reader = csv.reader(file)
        num_rows = sum(1 for row in csv_reader)
        return num_rows

def calculate_num_timesteps_excel(file_path):
    excel_data = pd.read_excel(file_path)
    num_rows = excel_data.shape[0]
    return num_rows

def calculate_num_timesteps_excel(file_path):
    excel_data = pd.read_excel(file_path)
    num_rows = excel_data.shape[0]
    return num_rows

def calculate_num_timesteps_word(file_path):
    doc = Document(file_path)
    num_paragraphs = len(doc.paragraphs)
    return num_paragraphs


def extract_text_from_pdf(file):
    text = ""
    pdf_reader = PyPDF2.PdfFileReader(file)
    num_pages = pdf_reader.getNumPages()
    for page_number in range(num_pages):
        page = pdf_reader.getPage(page_number)
        text += page.extractText()
    return text

def extract_text_from_docx(file):
    text = ""
    doc = Document(file)
    for paragraph in doc.paragraphs:
        text += paragraph.text + " "
    return text

def get_data_type(input_data):
    # Cria uma instância do objeto Magic
    
    content_type, encoding = mimetypes.guess_type(input_data)

    if content_type.startswith('image'):
        return 'Imagens', 'image'
    elif content_type.startswith('video'):
        return 'Vídeos', 'video'
    elif content_type.startswith('audio'):
        return 'Áudio', 'audio'
    elif content_type.startswith('text'):
        return 'Texto', 'text'
    else:
        if content_type == 'application/pdf':
            return 'Arquivos', 'pdf'
        elif content_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            return 'Arquivos', 'xls'
        elif content_type == 'text/csv':
            return 'Arquivos', 'csv'
        elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return 'Arquivos', 'doc'
        else:   
            return 'Arquivos', 'file'
        
def get_numtime_steps(input_data, data_subtype):
    if data_subtype == 'audio':
        # Calcular o num_timesteps específico para dados de áudio
        return calculate_num_timesteps_audio(input_data)
    elif data_subtype == 'video':
        # Calcular o num_timesteps específico para dados de vídeo
        return calculate_num_timesteps_video(input_data)
    elif data_subtype == 'image':
        # Calcular o num_timesteps específico para dados de imagem
        return calculate_num_timesteps_image(input_data)
    elif data_subtype == 'text':
        # Calcular o num_timesteps específico para dados de texto
        return calculate_num_timesteps_text(input_data)
    elif data_subtype == 'pdf':
        # Calcular o num_timesteps específico para outros tipos de arquivo
        return calculate_num_timesteps_pdf(input_data)
    elif data_subtype == 'xls':
         # Calcular o num_timesteps específico para outros tipos de arquivo
        return calculate_num_timesteps_excel(input_data)
    elif data_subtype == 'csv':
        # Calcular o num_timesteps específico para outros tipos de arquivo
        return calculate_num_timesteps_csv(input_data)
    elif data_subtype == 'doc':
        # Calcular o num_timesteps específico para outros tipos de arquivo
        return calculate_num_timesteps_word(input_data)
    elif data_subtype == 'file':
        # Calcular o num_timesteps específico para outros tipos de arquivo
        return calculate_num_timesteps_file(input_data)

def get_extracted_data(input_data, subtype):
    if subtype == 'pdf':
        return extract_text_from_pdf(input_data)
    elif subtype == 'doc':
        return extract_text_from_docx(input_data)
    else:
        return input_data